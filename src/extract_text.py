"""
Text Extraction Pipeline for Healthcare Consent Documents
Extracts raw text from PDF and DOCX files while preserving structure.
"""

import os
from pathlib import Path
from typing import Dict, List
import pdfplumber
from docx import Document
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extracts text from PDF and DOCX healthcare documents."""
    
    def __init__(self, raw_dir: str = "data/raw", output_dir: str = "data/processed/text"):
        """
        Initialize the text extractor.
        
        Args:
            raw_dir: Directory containing raw PDF and DOCX files
            output_dir: Directory to save extracted text files
        """
        self.raw_dir = Path(raw_dir)
        self.output_dir = Path(output_dir)
        self.pdf_dir = self.raw_dir / "pdf"
        self.docx_dir = self.raw_dir / "docx"
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Statistics
        self.stats = {
            "pdf_processed": 0,
            "docx_processed": 0,
            "pdf_failed": 0,
            "docx_failed": 0,
            "files_filtered_in": 0
        }
    
    def normalize_text(self, text: str) -> str:
        """
        Normalize text by removing trailing whitespace while preserving structure.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Normalized text with clean line breaks
        """
        # Remove trailing whitespace from each line
        # Preserves headings, section structure, and intentional line breaks
        normalized = "\n".join(line.rstrip() for line in text.splitlines())
        return normalized
    
    def extract_from_pdf(self, pdf_path: Path) -> str:
        """
        Extract text from a PDF file using pdfplumber.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            text_parts = []
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text from page (use empty string if extraction fails)
                    # This preserves page order and prevents silent content loss
                    page_text = page.extract_text() or ""
                    
                    if not page_text.strip():
                        # Log warning if page appears empty
                        logger.warning(f"Page {page_num} in {pdf_path.name} appears empty or failed to extract")
                        page_text = f"[PAGE {page_num} - NO TEXT EXTRACTED]"
                    
                    # Add page separator for multi-page documents
                    if page_num > 1:
                        text_parts.append(f"\n--- PAGE {page_num} ---\n")
                    text_parts.append(page_text)
            
            raw_text = "\n".join(text_parts)
            
            # Normalize line breaks
            return self.normalize_text(raw_text)
        
        except Exception as e:
            logger.error(f"Failed to extract from PDF {pdf_path.name}: {str(e)}")
            raise
    
    def extract_from_docx(self, docx_path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables (important for forms)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            raw_text = "\n".join(text_parts)
            
            # Normalize line breaks
            return self.normalize_text(raw_text)
        
        except Exception as e:
            logger.error(f"Failed to extract from DOCX {docx_path.name}: {str(e)}")
            raise
    
    def should_process_file(self, filename: str, consent_only: bool = True) -> bool:
        """
        Determine if a file should be processed based on naming.
        
        Args:
            filename: Name of the file
            consent_only: If True, only process files with 'consent' in name
            
        Returns:
            True if file should be processed
        """
        if consent_only:
            # Process files with 'consent' or 'intake' in the name
            keywords = ['consent', 'intake']
            return any(keyword in filename.lower() for keyword in keywords)
        return True
    
    def save_text(self, text: str, original_filename: str) -> Path:
        """
        Save extracted text to a file with encoding safety.
        
        Args:
            text: Extracted text content
            original_filename: Original file name (will be converted to .txt)
            
        Returns:
            Path to saved text file
        """
        # Create output filename (replace extension with .txt)
        output_filename = Path(original_filename).stem + ".txt"
        output_path = self.output_dir / output_filename
        
        # Save text with encoding error handling for broken PDF encodings
        with open(output_path, 'w', encoding='utf-8', errors='ignore') as f:
            f.write(text)
        
        return output_path
    
    def process_directory(self, directory: Path, file_type: str, consent_only: bool = True) -> List[Dict]:
        """
        Process all files in a directory.
        
        Args:
            directory: Directory containing files
            file_type: 'pdf' or 'docx'
            consent_only: If True, only process consent/intake files
            
        Returns:
            List of dictionaries with processing results
        """
        results = []
        
        if not directory.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return results
        
        # Get all files with the specified extension
        pattern = f"*.{file_type}"
        files = list(directory.glob(pattern))
        
        logger.info(f"Found {len(files)} {file_type.upper()} files in {directory}")
        
        for file_path in files:
            # Check if we should process this file
            if not self.should_process_file(file_path.name, consent_only):
                logger.info(f"Skipping {file_path.name} (not a consent/intake file)")
                continue
            
            if consent_only:
                self.stats["files_filtered_in"] += 1
            
            logger.info(f"Processing: {file_path.name}")
            
            try:
                # Extract text based on file type
                if file_type == 'pdf':
                    text = self.extract_from_pdf(file_path)
                    self.stats["pdf_processed"] += 1
                elif file_type == 'docx':
                    text = self.extract_from_docx(file_path)
                    self.stats["docx_processed"] += 1
                else:
                    raise ValueError(f"Unsupported file type: {file_type}")
                
                # Save extracted text
                output_path = self.save_text(text, file_path.name)
                
                # Record result
                result = {
                    "source_file": str(file_path),
                    "output_file": str(output_path),
                    "status": "success",
                    "char_count": len(text),
                    "line_count": text.count('\n') + 1
                }
                results.append(result)
                
                logger.info(f"✓ Extracted {len(text)} characters to {output_path.name}")
            
            except Exception as e:
                # Record failure
                if file_type == 'pdf':
                    self.stats["pdf_failed"] += 1
                else:
                    self.stats["docx_failed"] += 1
                
                result = {
                    "source_file": str(file_path),
                    "output_file": None,
                    "status": "failed",
                    "error": str(e)
                }
                results.append(result)
                
                logger.error(f"✗ Failed to process {file_path.name}")
        
        return results
    
    def extract_all(self, consent_only: bool = True) -> Dict:
        """
        Extract text from all PDF and DOCX files.
        
        Args:
            consent_only: If True, only process files with 'consent'/'intake' in name
            
        Returns:
            Dictionary with processing results and statistics
        """
        logger.info("="*60)
        logger.info("Starting Text Extraction Pipeline")
        logger.info("="*60)
        
        all_results = {
            "pdf_results": [],
            "docx_results": []
        }
        
        # Process PDFs
        logger.info("\n📄 Processing PDF files...")
        all_results["pdf_results"] = self.process_directory(
            self.pdf_dir, 'pdf', consent_only
        )
        
        # Process DOCX files
        logger.info("\n📝 Processing DOCX files...")
        all_results["docx_results"] = self.process_directory(
            self.docx_dir, 'docx', consent_only
        )
        
        # Print statistics
        logger.info("\n" + "="*60)
        logger.info("Extraction Complete - Statistics:")
        logger.info("="*60)
        logger.info(f"PDFs processed: {self.stats['pdf_processed']}")
        logger.info(f"PDFs failed: {self.stats['pdf_failed']}")
        logger.info(f"DOCX processed: {self.stats['docx_processed']}")
        logger.info(f"DOCX failed: {self.stats['docx_failed']}")
        logger.info(f"Total consent/intake files processed: {self.stats['files_filtered_in']}")
        logger.info(f"Output directory: {self.output_dir}")
        logger.info("="*60)
        
        return all_results


def main():
    """Main execution function."""
    # Initialize extractor
    extractor = TextExtractor()
    
    # Extract from all files (consent/intake only by default)
    results = extractor.extract_all(consent_only=True)
    
    # You can set consent_only=False to process all files
    # results = extractor.extract_all(consent_only=False)


if __name__ == "__main__":
    main()